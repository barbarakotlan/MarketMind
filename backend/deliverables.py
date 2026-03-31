from __future__ import annotations

import json
import uuid
from io import BytesIO
from typing import Any, Dict, List, Optional

import yfinance as yf
from sqlalchemy import delete, func, select
from sqlalchemy.orm import Session

from models import ensemble_predict, create_dataset
from openrouter_client import DEFAULT_OPENROUTER_MODEL, create_structured_completion
from user_state_store import (
    Deliverable,
    DeliverableAssumption,
    DeliverableLink,
    DeliverableMemo,
    DeliverablePreflight,
    DeliverableReview,
    load_notifications,
    load_portfolio,
    load_watchlist,
    utcnow,
)


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DEFAULT_TEMPLATE_KEY = "investment_thesis_memo"
DELIVERABLE_STATUSES = {"draft", "active", "monitoring", "invalidated", "complete", "archived"}
REVIEW_TYPES = {"checkpoint", "closeout"}
SOURCE_TYPES = {"user", "marketmind", "ai"}
PRELIGHT_STATUS_ORDER = {"red": 0, "yellow": 1, "green": 2}

MEMO_JSON_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "executive_summary": {"type": "string"},
        "investment_thesis": {"type": "string"},
        "supporting_evidence": {"type": "array", "items": {"type": "string"}},
        "key_assumptions": {"type": "array", "items": {"type": "string"}},
        "risks": {"type": "array", "items": {"type": "string"}},
        "invalidation_conditions": {"type": "array", "items": {"type": "string"}},
        "catalysts": {"type": "array", "items": {"type": "string"}},
        "signals_and_market_context": {"type": "array", "items": {"type": "string"}},
        "linked_positioning": {"type": "string"},
        "what_would_change_my_mind": {"type": "string"},
        "conclusion": {"type": "string"},
    },
    "required": [
        "executive_summary",
        "investment_thesis",
        "supporting_evidence",
        "key_assumptions",
        "risks",
        "invalidation_conditions",
        "catalysts",
        "signals_and_market_context",
        "linked_positioning",
        "what_would_change_my_mind",
        "conclusion",
    ],
}

MEMO_SECTION_TITLES = {
    "executive_summary": "Executive Summary",
    "investment_thesis": "Investment Thesis",
    "supporting_evidence": "Supporting Evidence",
    "key_assumptions": "Key Assumptions",
    "risks": "Risks",
    "invalidation_conditions": "Invalidation Conditions",
    "catalysts": "Catalysts",
    "signals_and_market_context": "Signals and Market Context",
    "linked_positioning": "Linked Positioning",
    "what_would_change_my_mind": "What Would Change My Mind",
    "conclusion": "Conclusion",
}


class DeliverableError(Exception):
    def __init__(self, message: str, *, status_code: int = 400, payload: Optional[Dict[str, Any]] = None):
        super().__init__(message)
        self.status_code = status_code
        self.payload = payload or {}


def _serialize_timestamp(value):
    return value.isoformat() if value else None


def _coerce_uuid(value: str):
    try:
        return uuid.UUID(str(value))
    except Exception as exc:
        raise DeliverableError("Invalid deliverable identifier", status_code=400) from exc


def _deliverable_summary(row: Deliverable, latest_memo_version: Optional[int] = None) -> Dict[str, Any]:
    return {
        "id": str(row.id),
        "templateKey": row.template_key,
        "ticker": row.ticker,
        "title": row.title,
        "status": row.status,
        "confidence": row.confidence,
        "timeHorizon": row.time_horizon,
        "memoAudience": row.memo_audience,
        "latestMemoVersion": latest_memo_version,
        "createdAt": _serialize_timestamp(row.created_at),
        "updatedAt": _serialize_timestamp(row.updated_at),
        "closedAt": _serialize_timestamp(row.closed_at),
    }


def _deliverable_detail(row: Deliverable) -> Dict[str, Any]:
    payload = _deliverable_summary(row)
    payload.update(
        {
            "thesisStatement": row.thesis_statement,
            "bullCase": row.bull_case,
            "bearCase": row.bear_case,
            "invalidationConditions": row.invalidation_conditions,
            "catalysts": row.catalysts,
        }
    )
    return payload


def _assumption_payload(row: DeliverableAssumption) -> Dict[str, Any]:
    return {
        "id": str(row.id),
        "label": row.label,
        "value": row.value,
        "reason": row.reason,
        "confidence": row.confidence,
        "sourceType": row.source_type,
        "sortOrder": row.sort_order,
        "createdAt": _serialize_timestamp(row.created_at),
        "updatedAt": _serialize_timestamp(row.updated_at),
    }


def _review_payload(row: DeliverableReview) -> Dict[str, Any]:
    return {
        "id": str(row.id),
        "reviewType": row.review_type,
        "summary": row.summary,
        "whatChanged": row.what_changed,
        "outcomeRating": row.outcome_rating,
        "createdAt": _serialize_timestamp(row.created_at),
    }


def _preflight_payload(row: DeliverablePreflight) -> Dict[str, Any]:
    return {
        "id": str(row.id),
        "inputVersion": row.input_version,
        "status": row.status,
        "requiredQuestions": list(row.required_questions_json or []),
        "assumptions": list(row.assumptions_json or []),
        "sources": list(row.sources_json or []),
        "blockingReasons": list(row.blocking_reasons_json or []),
        "createdAt": _serialize_timestamp(row.created_at),
    }


def _memo_payload(row: DeliverableMemo) -> Dict[str, Any]:
    return {
        "id": str(row.id),
        "version": row.version,
        "modelSlug": row.model_slug,
        "generationStatus": row.generation_status,
        "structuredContent": dict(row.structured_content_json or {}),
        "mimeType": row.mime_type,
        "hasArtifact": bool(row.docx_blob),
        "createdAt": _serialize_timestamp(row.created_at),
        "errorMessage": row.error_message,
    }


def _get_deliverable_row(session: Session, clerk_user_id: str, deliverable_id: str) -> Deliverable:
    deliverable_uuid = _coerce_uuid(deliverable_id)
    row = session.scalar(
        select(Deliverable).where(
            Deliverable.id == deliverable_uuid,
            Deliverable.clerk_user_id == clerk_user_id,
        )
    )
    if row is None:
        raise DeliverableError("Deliverable not found", status_code=404)
    return row


def list_deliverables(session: Session, clerk_user_id: str) -> List[Dict[str, Any]]:
    rows = session.scalars(
        select(Deliverable)
        .where(Deliverable.clerk_user_id == clerk_user_id)
        .order_by(Deliverable.updated_at.desc(), Deliverable.created_at.desc())
    ).all()
    if not rows:
        return []

    deliverable_ids = [row.id for row in rows]
    memo_rows = session.scalars(
        select(DeliverableMemo)
        .where(DeliverableMemo.deliverable_id.in_(deliverable_ids))
        .order_by(DeliverableMemo.deliverable_id.asc(), DeliverableMemo.version.desc())
    ).all()
    latest_memo_versions: Dict[str, int] = {}
    for memo in memo_rows:
        latest_memo_versions.setdefault(str(memo.deliverable_id), memo.version)

    return [
        _deliverable_summary(row, latest_memo_versions.get(str(row.id)))
        for row in rows
    ]


def create_deliverable(session: Session, clerk_user_id: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    now = utcnow()
    ticker = str(payload.get("ticker", "")).strip().upper()
    if not ticker:
        raise DeliverableError("Ticker is required", status_code=400)

    title = str(payload.get("title", "")).strip() or f"{ticker} Investment Thesis Memo"
    template_key = str(payload.get("templateKey") or DEFAULT_TEMPLATE_KEY)
    row = Deliverable(
        clerk_user_id=clerk_user_id,
        template_key=template_key,
        ticker=ticker,
        title=title,
        thesis_statement=str(payload.get("thesisStatement", "")).strip() or None,
        time_horizon=str(payload.get("timeHorizon", "")).strip() or None,
        bull_case=str(payload.get("bullCase", "")).strip() or None,
        bear_case=str(payload.get("bearCase", "")).strip() or None,
        invalidation_conditions=str(payload.get("invalidationConditions", "")).strip() or None,
        catalysts=str(payload.get("catalysts", "")).strip() or None,
        status=str(payload.get("status") or "draft"),
        confidence=str(payload.get("confidence", "")).strip() or None,
        memo_audience=str(payload.get("memoAudience") or "personal investment review"),
        created_at=now,
        updated_at=now,
    )
    if row.status not in DELIVERABLE_STATUSES:
        row.status = "draft"
    session.add(row)
    session.flush()
    return _deliverable_detail(row)


def update_deliverable(session: Session, clerk_user_id: str, deliverable_id: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    field_map = {
        "ticker": ("ticker", lambda value: str(value).strip().upper()),
        "title": ("title", lambda value: str(value).strip()),
        "thesisStatement": ("thesis_statement", lambda value: str(value).strip()),
        "timeHorizon": ("time_horizon", lambda value: str(value).strip()),
        "bullCase": ("bull_case", lambda value: str(value).strip()),
        "bearCase": ("bear_case", lambda value: str(value).strip()),
        "invalidationConditions": ("invalidation_conditions", lambda value: str(value).strip()),
        "catalysts": ("catalysts", lambda value: str(value).strip()),
        "status": ("status", lambda value: str(value).strip()),
        "confidence": ("confidence", lambda value: str(value).strip()),
        "memoAudience": ("memo_audience", lambda value: str(value).strip()),
    }

    for request_key, (attr_name, normalizer) in field_map.items():
        if request_key not in payload:
            continue
        normalized = normalizer(payload.get(request_key))
        setattr(row, attr_name, normalized or None)

    if row.status not in DELIVERABLE_STATUSES:
        row.status = "draft"
    row.closed_at = utcnow() if row.status in {"invalidated", "complete", "archived"} else None
    row.updated_at = utcnow()
    session.flush()
    return _deliverable_detail(row)


def replace_deliverable_assumptions(
    session: Session,
    clerk_user_id: str,
    deliverable_id: str,
    assumptions: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    session.execute(delete(DeliverableAssumption).where(DeliverableAssumption.deliverable_id == row.id))
    session.flush()

    now = utcnow()
    created_rows: List[DeliverableAssumption] = []
    for index, item in enumerate(assumptions or []):
        label = str(item.get("label", "")).strip()
        value = str(item.get("value", "")).strip()
        if not label or not value:
            continue
        assumption = DeliverableAssumption(
            deliverable_id=row.id,
            label=label,
            value=value,
            reason=str(item.get("reason", "")).strip() or None,
            confidence=str(item.get("confidence", "")).strip() or None,
            source_type=str(item.get("sourceType") or "user"),
            sort_order=index,
            created_at=now,
            updated_at=now,
        )
        if assumption.source_type not in SOURCE_TYPES:
            assumption.source_type = "user"
        session.add(assumption)
        created_rows.append(assumption)

    row.updated_at = now
    session.flush()
    return [_assumption_payload(item) for item in created_rows]


def add_deliverable_review(
    session: Session,
    clerk_user_id: str,
    deliverable_id: str,
    payload: Dict[str, Any],
) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    summary = str(payload.get("summary", "")).strip()
    if not summary:
        raise DeliverableError("Review summary is required", status_code=400)

    review = DeliverableReview(
        deliverable_id=row.id,
        review_type=str(payload.get("reviewType") or "checkpoint"),
        summary=summary,
        what_changed=str(payload.get("whatChanged", "")).strip() or None,
        outcome_rating=str(payload.get("outcomeRating", "")).strip() or None,
        created_at=utcnow(),
    )
    if review.review_type not in REVIEW_TYPES:
        review.review_type = "checkpoint"
    session.add(review)
    row.updated_at = utcnow()
    session.flush()
    return _review_payload(review)


def _get_assumptions(session: Session, deliverable_id) -> List[DeliverableAssumption]:
    return session.scalars(
        select(DeliverableAssumption)
        .where(DeliverableAssumption.deliverable_id == deliverable_id)
        .order_by(DeliverableAssumption.sort_order.asc(), DeliverableAssumption.created_at.asc())
    ).all()


def _get_reviews(session: Session, deliverable_id) -> List[DeliverableReview]:
    return session.scalars(
        select(DeliverableReview)
        .where(DeliverableReview.deliverable_id == deliverable_id)
        .order_by(DeliverableReview.created_at.desc())
    ).all()


def _get_latest_preflight(session: Session, deliverable_id) -> Optional[DeliverablePreflight]:
    return session.scalar(
        select(DeliverablePreflight)
        .where(DeliverablePreflight.deliverable_id == deliverable_id)
        .order_by(DeliverablePreflight.created_at.desc())
    )


def list_deliverable_memos(session: Session, clerk_user_id: str, deliverable_id: str) -> List[Dict[str, Any]]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    memos = session.scalars(
        select(DeliverableMemo)
        .where(DeliverableMemo.deliverable_id == row.id)
        .order_by(DeliverableMemo.version.desc())
    ).all()
    return [_memo_payload(item) for item in memos]


def _prediction_snapshot(ticker: str) -> Optional[Dict[str, Any]]:
    try:
        df = create_dataset(ticker, period="1y")
        if df.empty or len(df) < 30:
            return None
        ensemble_preds, model_breakdown = ensemble_predict(df, days_ahead=6)
        if ensemble_preds is None or len(ensemble_preds) == 0:
            return None
        recent_close = float(df["Close"].iloc[-1])
        model_first_predictions = [float(preds[0]) for preds in model_breakdown.values() if preds is not None and len(preds)]
        confidence = 85.0
        if len(model_first_predictions) > 1:
            mean_prediction = sum(model_first_predictions) / len(model_first_predictions)
            dispersion = sum(abs(pred - mean_prediction) for pred in model_first_predictions) / len(model_first_predictions)
            if recent_close:
                confidence = max(55.0, min(95.0, 95.0 - ((dispersion / recent_close) * 100)))
        return {
            "recentClose": round(recent_close, 2),
            "recentPredicted": round(float(ensemble_preds[0]), 2),
            "confidence": round(confidence, 1),
            "modelsUsed": list(model_breakdown.keys()),
            "predictions": [
                {"day": index + 1, "predictedClose": round(float(pred), 2)}
                for index, pred in enumerate(ensemble_preds[:3])
            ],
        }
    except Exception:
        return None


def _recent_news(ticker: str) -> List[Dict[str, Any]]:
    try:
        news_items = yf.Ticker(ticker).news or []
    except Exception:
        news_items = []
    formatted: List[Dict[str, Any]] = []
    for item in news_items[:5]:
        formatted.append(
            {
                "title": item.get("title"),
                "publisher": item.get("publisher"),
                "link": item.get("link"),
                "publishedAt": item.get("providerPublishTime"),
            }
        )
    return formatted


def _fundamentals_summary(ticker: str) -> Dict[str, Any]:
    try:
        info = yf.Ticker(ticker).info or {}
    except Exception:
        info = {}
    return {
        "companyName": info.get("longName"),
        "sector": info.get("sector"),
        "industry": info.get("industry"),
        "marketCap": info.get("marketCap"),
        "trailingPE": info.get("trailingPE"),
        "forwardPE": info.get("forwardPE"),
        "targetMeanPrice": info.get("targetMeanPrice"),
        "fiftyTwoWeekHigh": info.get("fiftyTwoWeekHigh"),
        "fiftyTwoWeekLow": info.get("fiftyTwoWeekLow"),
        "summary": info.get("longBusinessSummary"),
    }


def _sync_context_links(session: Session, deliverable_id, context: Dict[str, Any]) -> None:
    session.execute(delete(DeliverableLink).where(DeliverableLink.deliverable_id == deliverable_id))
    now = utcnow()
    links: List[DeliverableLink] = []

    if context.get("watchlistMembership"):
        links.append(
            DeliverableLink(
                deliverable_id=deliverable_id,
                link_type="watchlist",
                link_ref=context["ticker"],
                metadata_json={"source": "watchlist"},
                created_at=now,
            )
        )

    for alert in context.get("activeAlerts", []):
        links.append(
            DeliverableLink(
                deliverable_id=deliverable_id,
                link_type="alert",
                link_ref=str(alert.get("id") or alert.get("ticker")),
                metadata_json=alert,
                created_at=now,
            )
        )

    for trade in context.get("paperTradeHistory", []):
        links.append(
            DeliverableLink(
                deliverable_id=deliverable_id,
                link_type="paper_trade",
                link_ref=str(trade.get("timestamp") or trade.get("date") or trade.get("ticker")),
                metadata_json=trade,
                created_at=now,
            )
        )

    if context.get("predictionSnapshot"):
        links.append(
            DeliverableLink(
                deliverable_id=deliverable_id,
                link_type="prediction_snapshot",
                link_ref=context["ticker"],
                metadata_json=context["predictionSnapshot"],
                created_at=now,
            )
        )

    for link in links:
        session.add(link)


def build_deliverable_context(session: Session, clerk_user_id: str, deliverable_id: str) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    ticker = row.ticker
    watchlist = load_watchlist(session, clerk_user_id)
    notifications = load_notifications(session, clerk_user_id)
    portfolio = load_portfolio(session, clerk_user_id)

    active_alerts = [
        alert
        for alert in notifications.get("active", [])
        if str(alert.get("ticker", "")).upper() == ticker
    ]
    paper_trade_history = [
        trade
        for trade in portfolio.get("trade_history", [])[-50:]
        if str(trade.get("ticker", "")).upper() == ticker
    ][-10:]
    current_position = dict((portfolio.get("positions") or {}).get(ticker) or {})

    return {
        "ticker": ticker,
        "watchlistMembership": ticker in watchlist,
        "activeAlerts": active_alerts,
        "predictionSnapshot": _prediction_snapshot(ticker),
        "recentNews": _recent_news(ticker),
        "fundamentalsSummary": _fundamentals_summary(ticker),
        "paperTradeHistory": paper_trade_history,
        "currentPaperPosition": current_position,
    }


def _preflight_sources(context: Dict[str, Any]) -> List[Dict[str, Any]]:
    sources: List[Dict[str, Any]] = []
    if context.get("predictionSnapshot"):
        sources.append({"type": "prediction", "label": "Latest ensemble prediction snapshot"})
    if context.get("recentNews"):
        sources.append({"type": "news", "label": f"{len(context['recentNews'])} recent news items"})
    if context.get("fundamentalsSummary", {}).get("companyName"):
        sources.append({"type": "fundamentals", "label": "Fundamentals summary"})
    if context.get("paperTradeHistory"):
        sources.append({"type": "paper_trade", "label": f"{len(context['paperTradeHistory'])} paper trade events"})
    if context.get("activeAlerts"):
        sources.append({"type": "alerts", "label": f"{len(context['activeAlerts'])} active alerts"})
    return sources


def run_preflight(
    deliverable: Deliverable,
    assumptions: List[DeliverableAssumption],
    context: Dict[str, Any],
) -> Dict[str, Any]:
    required_questions: List[Dict[str, Any]] = []
    blocking_reasons: List[Dict[str, Any]] = []

    high_fields = {
        "ticker": deliverable.ticker,
        "title": deliverable.title,
        "thesisStatement": deliverable.thesis_statement,
        "timeHorizon": deliverable.time_horizon,
    }
    medium_fields = {
        "bullCase": deliverable.bull_case,
        "bearCase": deliverable.bear_case,
        "invalidationConditions": deliverable.invalidation_conditions,
    }
    low_fields = {
        "catalysts": deliverable.catalysts,
        "assumptions": assumptions,
        "linkedContext": _preflight_sources(context),
    }

    for field_name, value in high_fields.items():
        if not value:
            item = {
                "field": field_name,
                "severity": "high",
                "message": f"Add {field_name} before generating a memo.",
            }
            required_questions.append(item)
            blocking_reasons.append(item)

    for field_name, value in medium_fields.items():
        if not value:
            required_questions.append(
                {
                    "field": field_name,
                    "severity": "medium",
                    "message": f"Consider filling {field_name} to strengthen the memo.",
                }
            )

    if not low_fields["catalysts"]:
        required_questions.append(
            {
                "field": "catalysts",
                "severity": "low",
                "message": "Add catalysts or events to watch.",
            }
        )
    if not low_fields["assumptions"]:
        required_questions.append(
            {
                "field": "assumptions",
                "severity": "low",
                "message": "Add at least one explicit assumption.",
            }
        )
    if not low_fields["linkedContext"]:
        required_questions.append(
            {
                "field": "linkedContext",
                "severity": "low",
                "message": "The memo will be stronger once MarketMind context is available.",
            }
        )

    status = "green"
    if blocking_reasons:
        status = "red"
    elif required_questions:
        status = "yellow"

    return {
        "status": status,
        "requiredQuestions": required_questions,
        "assumptions": [
            {
                "label": assumption.label,
                "value": assumption.value,
                "reason": assumption.reason,
                "confidence": assumption.confidence,
                "sourceType": assumption.source_type,
            }
            for assumption in assumptions
        ],
        "sources": _preflight_sources(context),
        "blockingReasons": blocking_reasons,
    }


def create_deliverable_preflight(session: Session, clerk_user_id: str, deliverable_id: str) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    assumptions = _get_assumptions(session, row.id)
    context = build_deliverable_context(session, clerk_user_id, deliverable_id)
    payload = run_preflight(row, assumptions, context)
    input_version = int(
        (session.scalar(
            select(func.count()).select_from(DeliverablePreflight).where(DeliverablePreflight.deliverable_id == row.id)
        ) or 0)
        + 1
    )
    preflight = DeliverablePreflight(
        deliverable_id=row.id,
        input_version=input_version,
        status=payload["status"],
        required_questions_json=payload["requiredQuestions"],
        assumptions_json=payload["assumptions"],
        sources_json=payload["sources"],
        blocking_reasons_json=payload["blockingReasons"],
        created_at=utcnow(),
    )
    session.add(preflight)
    _sync_context_links(session, row.id, context)
    row.updated_at = utcnow()
    session.flush()
    return _preflight_payload(preflight)


def get_deliverable_detail(session: Session, clerk_user_id: str, deliverable_id: str) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    assumptions = _get_assumptions(session, row.id)
    reviews = _get_reviews(session, row.id)
    latest_preflight = _get_latest_preflight(session, row.id)
    memos = session.scalars(
        select(DeliverableMemo)
        .where(DeliverableMemo.deliverable_id == row.id)
        .order_by(DeliverableMemo.version.desc())
    ).all()
    latest_memo_version = memos[0].version if memos else None
    detail = _deliverable_detail(row)
    detail["latestMemoVersion"] = latest_memo_version
    return {
        "deliverable": detail,
        "assumptions": [_assumption_payload(item) for item in assumptions],
        "reviews": [_review_payload(item) for item in reviews],
        "latestPreflight": _preflight_payload(latest_preflight) if latest_preflight else None,
        "memoVersions": [_memo_payload(item) for item in memos],
        "linkedContext": build_deliverable_context(session, clerk_user_id, deliverable_id),
    }


def _next_memo_version(session: Session, deliverable_id) -> int:
    current_max = session.scalar(
        select(func.max(DeliverableMemo.version)).where(DeliverableMemo.deliverable_id == deliverable_id)
    )
    return int(current_max or 0) + 1


def _build_prompt_payload(
    deliverable: Dict[str, Any],
    assumptions: List[Dict[str, Any]],
    context: Dict[str, Any],
    latest_preflight: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    return {
        "template": DEFAULT_TEMPLATE_KEY,
        "instructions": [
            "Write a serious investment thesis memo using only the provided structured thesis data and MarketMind context.",
            "Do not invent facts, catalysts, or numeric claims not present in the provided context.",
            "If evidence is weak or missing, say so explicitly rather than filling gaps.",
            "Keep the tone analytical, clear, and useful for personal investment review.",
        ],
        "deliverable": deliverable,
        "assumptions": assumptions,
        "linkedContext": context,
        "latestPreflight": latest_preflight,
    }


def _build_openrouter_messages(prompt_payload: Dict[str, Any]) -> List[Dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "You are drafting an investment thesis memo for a market research application. "
                "Use only the structured thesis fields and context provided by the application. "
                "Return JSON only and follow the schema exactly."
            ),
        },
        {
            "role": "user",
            "content": (
                "Draft an investment thesis memo for the deliverable below.\n\n"
                f"{json.dumps(prompt_payload, indent=2)}"
            ),
        },
    ]


def _render_docx(deliverable: Dict[str, Any], structured_content: Dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to render memo artifacts") from exc

    document = Document()
    document.add_heading(deliverable.get("title") or "Investment Thesis Memo", 0)
    subtitle = f"{deliverable.get('ticker', '')} | {deliverable.get('status', 'draft').title()} | {deliverable.get('timeHorizon') or 'Unspecified horizon'}"
    document.add_paragraph(subtitle)

    for key, title in MEMO_SECTION_TITLES.items():
        value = structured_content.get(key)
        document.add_heading(title, level=1)
        if isinstance(value, list):
            if not value:
                document.add_paragraph("No content provided.")
            for item in value:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(value or "No content provided."))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_deliverable_memo(session: Session, clerk_user_id: str, deliverable_id: str) -> Dict[str, Any]:
    detail = get_deliverable_detail(session, clerk_user_id, deliverable_id)
    deliverable = detail["deliverable"]
    latest_preflight = detail["latestPreflight"] or create_deliverable_preflight(session, clerk_user_id, deliverable_id)
    if latest_preflight["status"] == "red":
        raise DeliverableError(
            "Preflight is red; complete the required sections before generating.",
            status_code=409,
            payload={"latestPreflight": latest_preflight},
        )

    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    assumptions = detail["assumptions"]
    context = detail["linkedContext"]
    _sync_context_links(session, row.id, context)

    version = _next_memo_version(session, row.id)
    prompt_payload = _build_prompt_payload(deliverable, assumptions, context, latest_preflight)
    messages = _build_openrouter_messages(prompt_payload)

    try:
        ai_result = create_structured_completion(
            messages=messages,
            json_schema=MEMO_JSON_SCHEMA,
            schema_name="investment_thesis_memo",
            model=DEFAULT_OPENROUTER_MODEL,
        )
        structured_content = ai_result["structured_content"]
        artifact = _render_docx(deliverable, structured_content)
        memo = DeliverableMemo(
            deliverable_id=row.id,
            version=version,
            model_slug=ai_result["model"],
            generation_status="completed",
            prompt_snapshot_json={"messages": messages},
            context_snapshot_json=context,
            structured_content_json=structured_content,
            docx_blob=artifact,
            mime_type=DOCX_MIME_TYPE,
            created_at=utcnow(),
            error_message=None,
        )
        session.add(memo)
        row.updated_at = utcnow()
        session.flush()
        return _memo_payload(memo)
    except Exception as exc:
        memo = DeliverableMemo(
            deliverable_id=row.id,
            version=version,
            model_slug=DEFAULT_OPENROUTER_MODEL,
            generation_status="failed",
            prompt_snapshot_json={"messages": messages},
            context_snapshot_json=context,
            structured_content_json={},
            docx_blob=None,
            mime_type=None,
            created_at=utcnow(),
            error_message=str(exc),
        )
        session.add(memo)
        row.updated_at = utcnow()
        session.flush()
        failed_payload = _memo_payload(memo)
        failed_payload["_statusCode"] = 502
        return failed_payload


def get_deliverable_memo_artifact(
    session: Session,
    clerk_user_id: str,
    deliverable_id: str,
    memo_id: str,
) -> Dict[str, Any]:
    row = _get_deliverable_row(session, clerk_user_id, deliverable_id)
    memo_uuid = _coerce_uuid(memo_id)
    memo = session.scalar(
        select(DeliverableMemo).where(
            DeliverableMemo.id == memo_uuid,
            DeliverableMemo.deliverable_id == row.id,
        )
    )
    if memo is None:
        raise DeliverableError("Memo version not found", status_code=404)
    if not memo.docx_blob:
        raise DeliverableError("Memo artifact is not available for this version", status_code=404)
    return {
        "filename": f"{row.ticker.lower()}-investment-thesis-memo-v{memo.version}.docx",
        "mimeType": memo.mime_type or DOCX_MIME_TYPE,
        "bytes": memo.docx_blob,
    }
